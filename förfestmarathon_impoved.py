import pandas as pd
import docx
from optimizer import TvisteOptimizer

# Constants
ADDRESS_COL = 7
TEL_COL = 8
HOST_NAME_COL = 3
TEMA_COL = 1
DEST_COL = 6


def summary(doc, groups_starter, groups_main, groups_desert, data, i):

    # Constants
    Guest1 = 3
    Guest2 = 4
    Guest3 = 5
    Spec_diet1 = 9
    Spec_diet2 = 10
    Spec_diet3 = 11
    Spec_diet4 = 12


    ## Förfest 1
    group_nr = groups_starter[i][0] - 1

    medlem1 = str(data.iloc[group_nr, Guest1])
    medlem2 = str(data.iloc[group_nr, Guest2])
    medlem3 = str(data.iloc[group_nr, Guest3])

    guest_1 = groups_starter[i][1] - 1
    guest_2 = groups_starter[i][2] - 1

    print("Förfest", group_nr, guest_1, guest_2)

    special_diet11 = str(data.iloc[guest_1, Spec_diet1])
    special_diet21 = str(data.iloc[guest_2, Spec_diet1])

    special_diet12 = str(data.iloc[guest_1, Spec_diet2])
    special_diet22 = str(data.iloc[guest_2, Spec_diet2])

    special_diet13 = str(data.iloc[guest_1, Spec_diet3])
    special_diet23 = str(data.iloc[guest_2, Spec_diet3])

    special_diet14 = str(data.iloc[guest_1, Spec_diet4])
    special_diet24 = str(data.iloc[guest_2, Spec_diet4])

    special_kost_line1 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)
    special_kost_line1 = special_kost_line1.replace('KOST3', special_diet13)
    special_kost_line1 = special_kost_line1.replace('KOST4', special_diet14)

    special_kost_line2 = 'Specialkost: KOST1 , KOST2 , KOST3  , KOST4'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)
    special_kost_line2 = special_kost_line2.replace('KOST3', special_diet23)
    special_kost_line2 = special_kost_line2.replace('KOST4', special_diet24)


    group_header = 'Grupp: Namn1, Namn2, Namn3'
    group_header = group_header.replace('Namn1', medlem1)
    group_header = group_header.replace('Namn2', medlem2)
    group_header = group_header.replace('Namn3', medlem3)

    doc.add_heading(group_header, 2)
    doc.add_paragraph('Ni ska styra kvällens första förfest klockan 17:45. Ta hänsyn till följande specialkost: ')
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")

    ##Förfest 2
    group_nr = groups_main[i][0] - 1

    medlem1 = str(data.iloc[group_nr, Guest1])
    medlem2 = str(data.iloc[group_nr, Guest2])
    medlem3 = str(data.iloc[group_nr, Guest3])

    guest_1 = groups_main[i][1] - 1
    guest_2 = groups_main[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, Spec_diet1])
    special_diet21 = str(data.iloc[guest_2, Spec_diet1])

    special_diet12 = str(data.iloc[guest_1, Spec_diet2])
    special_diet22 = str(data.iloc[guest_2, Spec_diet2])

    special_diet13 = str(data.iloc[guest_1, Spec_diet3])
    special_diet23 = str(data.iloc[guest_2, Spec_diet3])

    special_diet14 = str(data.iloc[guest_1, Spec_diet4])
    special_diet24 = str(data.iloc[guest_2, Spec_diet4])

    special_kost_line1 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)
    special_kost_line1 = special_kost_line1.replace('KOST3', special_diet13)
    special_kost_line1 = special_kost_line1.replace('KOST4', special_diet14)

    special_kost_line2 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)
    special_kost_line2 = special_kost_line2.replace('KOST3', special_diet23)
    special_kost_line2 = special_kost_line2.replace('KOST4', special_diet24)


    group_header = 'Grupp: Namn1, Namn2, Namn3'
    group_header = group_header.replace('Namn1', medlem1)
    group_header = group_header.replace('Namn2', medlem2)
    group_header = group_header.replace('Namn3', medlem3)

    doc.add_heading(group_header, 2)
    doc.add_paragraph('Ni ska styra kvällens andra förfest klockan 19:15. Ta hänsyn till följande specialkost: ')
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")

    ## Förfest 3
    group_nr = groups_desert[i][0] - 1

    medlem1 = str(data.iloc[group_nr, Guest1])
    medlem2 = str(data.iloc[group_nr, Guest2])
    medlem3 = str(data.iloc[group_nr, Guest3])

    guest_1 = groups_desert[i][1] - 1
    guest_2 = groups_desert[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, Spec_diet1])
    special_diet21 = str(data.iloc[guest_2, Spec_diet1])

    special_diet12 = str(data.iloc[guest_1, Spec_diet2])
    special_diet22 = str(data.iloc[guest_2, Spec_diet2])

    special_diet13 = str(data.iloc[guest_1, Spec_diet3])
    special_diet23 = str(data.iloc[guest_2, Spec_diet3])

    special_diet14 = str(data.iloc[guest_1, Spec_diet4])
    special_diet24 = str(data.iloc[guest_2, Spec_diet4])

    special_kost_line1 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)
    special_kost_line1 = special_kost_line1.replace('KOST3', special_diet13)
    special_kost_line1 = special_kost_line1.replace('KOST4', special_diet14)

    special_kost_line2 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)
    special_kost_line2 = special_kost_line2.replace('KOST3', special_diet23)
    special_kost_line2 = special_kost_line2.replace('KOST4', special_diet24)

    group_header = 'Grupp: Namn1, Namn2, Namn3'
    group_header = group_header.replace('Namn1', medlem1)
    group_header = group_header.replace('Namn2', medlem2)
    group_header = group_header.replace('Namn3', medlem3)

    doc.add_heading(group_header, 2)
    doc.add_paragraph('Ni ska styra kvällens sista förfest klockan 20:45. Ta hänsyn till följande specialkost: ')
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")


def get_group_info(dish, groups_starter, groups_main, groups_dessert, data, i):
    if dish == 'starter':
        group_nr = groups_starter[i][0] - 1
        main_destination_group = next((g[0] for g in groups_main if groups_starter[i][0] in g), None) - 1
        dessert_destination_group = next((g[0] for g in groups_dessert if groups_starter[i][0] in g), None) - 1
        host_1_nr, host_2_nr = main_destination_group, dessert_destination_group
    elif dish == 'main':
        group_nr = groups_main[i][0] - 1
        starter_destination_group = next((g[0] for g in groups_starter if groups_main[i][0] in g), None) - 1
        dessert_destination_group = next((g[0] for g in groups_dessert if groups_main[i][0] in g), None) - 1
        host_1_nr, host_2_nr = starter_destination_group, dessert_destination_group
    elif dish == 'dessert':
        group_nr = groups_dessert[i][0] - 1
        starter_destination_group = next((g[0] for g in groups_starter if groups_dessert[i][0] in g), None) - 1
        main_destination_group = next((g[0] for g in groups_main if groups_dessert[i][0] in g), None) - 1
        host_1_nr, host_2_nr = starter_destination_group, main_destination_group

    # Fetch and return group-related data
    country = data.iloc[group_nr, TEMA_COL]
    destination_1 = str(data.iloc[host_1_nr, DEST_COL])
    address_1 = str(data.iloc[host_1_nr, ADDRESS_COL])
    tel_1 = str(data.iloc[host_1_nr, TEL_COL])
    host_1_name = str(data.iloc[host_1_nr, HOST_NAME_COL])
    tema1 = str(data.iloc[host_1_nr, TEMA_COL])

    destination_2 = str(data.iloc[host_2_nr, DEST_COL])
    address_2 = str(data.iloc[host_2_nr, ADDRESS_COL])
    tel_2 = str(data.iloc[host_2_nr, TEL_COL])
    host_2_name = str(data.iloc[host_2_nr, HOST_NAME_COL])
    tema2 = str(data.iloc[host_2_nr, TEMA_COL])

    return country, destination_1, address_1, tel_1, host_1_name, tema1, destination_2, address_2, tel_2, host_2_name, tema2


def construct_base_doc_info( *info):


    # Första destinationen
    host_name1 = info[3]
    address_line1 = 'Adress: {}'.format(info[1])
    tema_line1 = 'Tema: {}'.format(info[4])
    tel_line1 = 'Telefonnummer: {} {}'.format(host_name1, info[2])  # Inkluderar både namn och telefonnummer

    # Andra destinationen
    host_name2 = info[8]
    address_line2 = 'Adress: {}'.format(info[6])
    tema_line2 = 'Tema: {}'.format(info[9])
    tel_line2 = 'Telefonnummer: {} {}'.format(host_name2, info[7])  # Inkluderar både namn och telefonnummer



    return {
        'address_line1': address_line1,
        'tema_line1': tema_line1,
        'tel_line1': tel_line1,
        'address_line2': address_line2,
        'tema_line2': tema_line2,
        'tel_line2': tel_line2,
    }


def create_fest_documentation(doc, dish, body_texts, doc_elements, country):
    # Konfiguration för att bestämma vilka paragrafer som ska läggas till
    config = {
        'starter': {2: ['tema_line1', 'address_line1', 'tel_line1'],
                    3: ['tema_line2', 'address_line2', 'tel_line2']},
        'main':    {1: ['tema_line1', 'address_line1', 'tel_line1'],
                    3: ['tema_line2', 'address_line2', 'tel_line2']},
        'dessert': {1: ['tema_line1', 'address_line1', 'tel_line1'],
                    2: ['tema_line1', 'address_line1', 'tel_line1'],
                    3: ['tema_line2', 'address_line2', 'tel_line2']}
    }

    # Hjälpfunktion för att lägga till paragrafer baserat på element från konfigurationen
    def add_paragraphs(doc, elements):
        for element in elements:
            if element in doc_elements:
                doc.add_paragraph(doc_elements[element])

    group_header = 'Förfestmarathon 2024 -  Grupp: {}'.format(country)
    doc.add_heading(group_header, 2)
    # Loopa genom varje text och dess index
    for index, body_text in enumerate(body_texts, start=1):
        doc.add_heading(f'Förfest {index}', 2)
        doc.add_paragraph(body_text)

        # Använd konfigurationen för att lägga till paragrafer baserade på rätt 'dish' och 'index'
        if index in config.get(dish, {}):
            add_paragraphs(doc, config[dish][index])
        doc.add_paragraph(" ")


    # Lägg till avslutande rubrik och sidbrytning
    doc.add_heading("22:00 Location: Rouge Nightclub", 2)
    doc.add_page_break()



def starter_documentation(doc, groups_starter, groups_main, groups_desert, data, i):
    dish = 'starter'
    country, *info = get_group_info(dish, groups_starter, groups_main, groups_desert, data, i)
    doc_elements = construct_base_doc_info(*info)

    body_texts = [
        'Resan tar nu sin början och det är dags för er att styra förfest. Skynda er hem, gästerna anländer 17:45',
        'Har ni koll på klockan? Nästa förfest drar igång 19:15. Ni ska nu ta er vidare till {}, Drick upp glasen och tacka för ett bra krök.'.format(info[1]),
        'Resan närmar sig sitt slut och det är dags för kvällens sista förfest. Förfesten drar igång 20:45. Skynda er till {} för att svalka era strupar innan gemensam utgång.'.format(info[2])
    ]
    create_fest_documentation(doc, dish, body_texts, doc_elements, country)

def main_course_documentation(doc, groups_starter, groups_main, groups_desert, data, i):
    dish = 'main'
    country, *info = get_group_info(dish, groups_starter, groups_main, groups_desert, data, i)
    doc_elements = construct_base_doc_info( *info)

    body_texts = [
        'Resan tar nu sin början. Skynda er till {} för kvällens första förfest. Festen drar igång 17:45'.format(info[1]),
        'Det är dags för er att styra förfest. Skynda er hem, gästerna anländer 19:15',
        'Resan närmar sig sitt slut och det är dags för kvällens sista förfest. Förfesten drar igång 20:45. Skynda er till {} för att svalka era strupar innan gemensam utgång.'.format(info[2])
    ]
    create_fest_documentation(doc, dish, body_texts, doc_elements, country)

def dessert_documentation(doc, groups_starter, groups_main, groups_desert, data, i):
    dish = 'dessert'
    country, *info = get_group_info(dish, groups_starter, groups_main, groups_desert, data, i)
    doc_elements = construct_base_doc_info(*info)

    body_texts = [
        'Resan tar nu sin början. Skynda er till {} för kvällens första förfest. Festen drar igång 17:45'.format(info[1]),
        'Har ni koll på klockan? Nästa förfest drar igång 19:15. Ni ska nu ta er vidare till {}, Drick upp glasen och tacka för ett bra krök.'.format(info[2]),
        'Resan närmar sig sitt slut och det är dags för kvällens grand finale. Skynda er hem och välkomna gästerna. Efter det möts vi för gemensam utgång 22:00.'
    ]
    create_fest_documentation(doc, dish, body_texts, doc_elements, country)


def grouping(data):
    if len(data) % 3 == 1:
        diff = 1
        data.drop(len(data) - 1, axis=0, inplace=True)
        print("Grupperingen går ej ihop.", diff, "st grupp ej inkluderad. Gruppen motsvarar den sista i svarsfilen")
    elif len(data) % 3 == 2:
        diff = 2
        data.drop([len(data) - 2, len(data) - 1], axis=0, inplace=True)
        print("Grupperingen går ej ihop.", diff, "st grupper ej inkluderad. Grupperna motsvarar de 2"
                                                 " sista i svarsfilen")
    else:
        print("Grupperingen går ihop, antal grupper: ", len(data) )

    groups_starter = []
    groups_main = []
    groups_desert = []
    for i in range(1, len(data), 3):
        groups_starter.append((i, i + 1, i + 2))

    for i in range(2, len(data), 3):
        if i + 4 == len(data):
            groups_main.append((i, i + 2, 3))
            groups_main.append((i + 3, 1, 6))
            break
        else:
            groups_main.append((i, i + 2, i + 7))

    for i in range(3, len(data), 3):
        if i + 3 == len(data):
            groups_desert.append((i, i + 1, 2))
            groups_desert.append((i + 3, 1, 5))
            break
        else:
            groups_desert.append((i, i + 1, i + 5))

    return groups_starter, groups_main, groups_desert


def main():

    """
    # option to lock groups together, or to a specific dish etc
    initial_locked = []

    # This is a optimizer that i based on brute forcing, optimizes grouping so that "Efterfest på tviste" and that
    # groups from specifik programs meet peopble from athoer programs
    optimizer = TvisteOptimizer('Förfestmarathon 2024.xlsx', 'Vilket område kommer ni vara på?', 'Viket program pluggar ni?', initial_locked)

    # Since its based on randomness, raise the iterations for the final run. I suggest 1000 iterations for a optimal solution
    # Saves the optimal solution to final_optimized.xlsx if you want a overwiev of the results
    finaldata = optimizer.optimize(iterations=100)
    print("The final optimized data has been saved in:", finaldata)
    """
    # Reads the optimized grouping
    df = pd.read_excel(r'final_optimized.xlsx', header=0)

    #Grouping based on a fixed mathematical formula
    groups_starter, groups_main, groups_desert = grouping(df)
    print('grops starter',groups_starter)
    print('groups main', groups_main)
    print('groups dessert', groups_desert)



    #Creates the documents "Breven"
    doc1 = docx.Document()
    doc2 = docx.Document()
    doc3 = docx.Document()
    for i in range(0, len(groups_starter)):
        starter_documentation(doc1, groups_starter, groups_main, groups_desert, df, i)
        main_course_documentation(doc2, groups_starter, groups_main, groups_desert, df, i)
        dessert_documentation(doc3, groups_starter, groups_main, groups_desert, df, i)
    doc1.save('outputs/Förrätter.docx')
    doc2.save('outputs/Varmrätter.docx')
    doc3.save('outputs/Efterrätter.docx')


    # Skapa summary dokument för publicering innan event. Innehållande vilken grupp som har vilken rätt och vilken specialkost de bör ta hänsyn till
    doc = docx.Document()
    for k in range(0, len(groups_starter)):
        summary(doc, groups_starter, groups_main, groups_desert, df, k)
    doc.save('outputs/Summary.docx')





if __name__ == '__main__':
    main()