""""
    ==========================================
     Title:  Cykelfest
     Author: Max Apelquist
     Date:   19 Sep 2022
    ==========================================
"""


import pandas as pd
import docx


def starter_documentation(doc, groups_starter, groups_main, groups_desert, data, i):

    starter_group_nr = groups_starter[i][0]
    for var in range(0, len(groups_main)):
        if starter_group_nr in groups_main[var]:
            main_destination_group = groups_main[var][0]

    for var1 in range(0, len(groups_desert)):
        if starter_group_nr in groups_desert[var1]:
            desert_destination_group = groups_desert[var1][0]

    group_nr = groups_starter[i][0] - 1
    country = data.iloc[group_nr, 1]
    name = data.iloc[group_nr, 2]

    host_1_nr = main_destination_group - 1
    destination_1 = data.iloc[host_1_nr, 4]
    address_1 = data.iloc[host_1_nr, 5]
    tel_1 = data.iloc[host_1_nr, 6]
    host_1_name = data.iloc[host_1_nr, 2]

    host_2_nr = desert_destination_group - 1
    destination_2 = data.iloc[host_2_nr, 4]
    address_2 = data.iloc[host_2_nr, 5]
    tel_2 = data.iloc[host_2_nr, 6]
    host_2_name = data.iloc[host_2_nr, 2]

    group_header = 'Grupp LAND, NAMN'
    group_header = group_header.replace('LAND', country)
    group_header = group_header.replace('NAMN', name)

    address_line = 'Adress: ADRESS'
    tel_line = 'Telefonnummer: NAME TEL'
    address_line = address_line.replace('ADRESS', address_1)
    tel_line = tel_line.replace('NAME', host_1_name)
    tel_line = tel_line.replace('TEL', tel_1)

    address_line2 = 'Adress: ADRESS'
    tel_line2 = 'Telefonnummer: NAME TEL'
    address_line2 = address_line2.replace('ADRESS', address_2)
    tel_line2 = tel_line2.replace('NAME', host_2_name)
    tel_line2 = tel_line2.replace('TEL', tel_2)

    guest_1 = groups_starter[i][1] - 1
    guest_2 = groups_starter[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, 7])
    special_diet21 = str(data.iloc[guest_2, 7])
    special_diet12 = str(data.iloc[guest_1, 8])
    special_diet22 = str(data.iloc[guest_2, 8])

    special_kost_line1 = 'Specialkost: KOST1 KOST2'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)

    special_kost_line2 = 'Specialkost: KOST1 KOST2'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)

    body_text_1 = 'Redo för något exotiskt, det är nu dags för varmrätt. Ni ska nu ta cykeln vidare till OMRÅDE1,' \
                ' så drick upp glasen och tacka för förrätten.'
    body_text_2 = 'Resan närmar sig sitt slut och det är dags för kvällens grand finale. Skynda er till OMRÅDE2 för att ' \
                'avnjuta en efterrätt. Efter det möts vi för gemensam utgång, ses där 22:00.'
    body_text_1 = body_text_1.replace('OMRÅDE1', destination_1)
    body_text_2 = body_text_2.replace('OMRÅDE2', destination_2)

    doc.add_heading(group_header, 2)
    doc.add_heading('CykelfIEzten 2021', 2)
    doc.add_paragraph('Resan tar nu sin början och det är dags för förrätt. Skynda er hem, för snart anländer gästerna')
    doc.add_paragraph(" ")
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")

    doc.add_heading('CykelfIEzten 2021', 2)
    doc.add_paragraph(body_text_1)
    doc.add_paragraph(address_line)
    doc.add_paragraph(tel_line)
    doc.add_paragraph(" ")

    doc.add_heading('CykelfIEzten 2021', 2)
    doc.add_paragraph(body_text_2)
    doc.add_paragraph(address_line2)
    doc.add_paragraph(tel_line2)
    doc.add_paragraph(" ")

    doc.add_page_break()


def main_course_documentation(doc1, groups_starter, groups_main, groups_desert, data, i):
    main_group_nr = groups_main[i][0]
    for var in range(0, len(groups_starter)):
        if main_group_nr in groups_starter[var]:
            starter_destination_group = groups_starter[var][0]

    for var1 in range(0, len(groups_desert)):
        if main_group_nr in groups_desert[var1]:
            desert_destination_group = groups_desert[var1][0]

    group_nr = groups_main[i][0] - 1
    country = data.iloc[group_nr, 1]
    name = data.iloc[group_nr, 2]

    host_1_nr = starter_destination_group - 1
    destination_1 = data.iloc[host_1_nr, 4]
    address_1 = data.iloc[host_1_nr, 5]
    tel_1 = data.iloc[host_1_nr, 6]
    host_1_name = data.iloc[host_1_nr, 2]

    host_2_nr = desert_destination_group - 1
    destination_2 = data.iloc[host_2_nr, 4]
    address_2 = data.iloc[host_2_nr, 5]
    tel_2 = data.iloc[host_2_nr, 6]
    host_2_name = data.iloc[host_2_nr, 2]

    guest_1 = groups_main[i][1] - 1
    guest_2 = groups_main[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, 7])
    special_diet21 = str(data.iloc[guest_2, 7])
    special_diet12 = str(data.iloc[guest_1, 8])
    special_diet22 = str(data.iloc[guest_2, 8])

    special_diet_line1 = 'Specialkost: KOST1 KOST2'
    special_diet_line1 = special_diet_line1.replace('KOST1', special_diet11)
    special_diet_line1 = special_diet_line1.replace('KOST2', special_diet12)

    special_diet_line2 = 'Specialkost: KOST1 KOST2'
    special_diet_line2 = special_diet_line2.replace('KOST1', special_diet21)
    special_diet_line2 = special_diet_line2.replace('KOST2', special_diet22)

    group_header = 'Grupp LAND, NAMN'
    group_header = group_header.replace('LAND', country)
    group_header = group_header.replace('NAMN', name)

    adress_line = 'Adress: ADRESS'
    tel_line = 'Telefonnummer: NAME TEL'
    adress_line = adress_line.replace('ADRESS', address_1)
    tel_line = tel_line.replace('NAME', host_1_name)
    tel_line = tel_line.replace('TEL', tel_1)

    adress_line2 = 'Adress: ADRESS'
    tel_line2 = 'Telefonnummer: NAME TEL'
    adress_line2 = adress_line2.replace('ADRESS', address_2)
    tel_line2 = tel_line2.replace('NAME', host_2_name)
    tel_line2 = tel_line2.replace('TEL', tel_2)

    body_text_1 = 'Resan tar nu sin början och det är nu dags för förrätt. Vässa smaklökarna och bege er till er första '\
                'destination: OMRÅDE1.'
    body_text_2 = 'Resan närmar sig sitt slut och det är dags för kvällens grand finale. Skynda er till OMRÅDE2 för att' \
                ' avnjuta en efterrätt. Efter det möts vi för gemensam utgång, ses där 22:00.'
    body_text_1 = body_text_1.replace('OMRÅDE1', destination_1)
    body_text_2 = body_text_2.replace('OMRÅDE2', destination_2)

    doc1.add_heading(group_header, 2)

    doc1.add_heading('CykelfIEzten 2021', 2)
    doc1.add_paragraph(body_text_1)
    doc1.add_paragraph(adress_line)
    doc1.add_paragraph(tel_line)
    doc1.add_paragraph(" ")

    doc1.add_heading('CykelfIEzten 2021', 2)
    doc1.add_paragraph(
        'Nu är det dags för varmrätt. Drick upp glasen, tacka för förrätten och skynda er hem! Gästerna kommer snart. ')
    doc1.add_paragraph(" ")
    doc1.add_paragraph(special_diet_line1)
    doc1.add_paragraph(special_diet_line2)
    doc1.add_paragraph(" ")

    doc1.add_heading('CykelfIEzten 2021', 2)
    doc1.add_paragraph(body_text_2)
    doc1.add_paragraph(adress_line2)
    doc1.add_paragraph(tel_line2)
    doc1.add_paragraph(" ")

    doc1.add_page_break()


def desert_documentation(doc2, groups_starter, groups_main, groups_desert, data, i):

    desert_group_nr = groups_desert[i][0]
    for var in range(0, len(groups_starter)):
        if desert_group_nr in groups_starter[var]:
            starter_destination_group = groups_starter[var][0]

    for var1 in range(0, len(groups_main)):
        if desert_group_nr in groups_main[var1]:
            main_destination_group = groups_main[var1][0]

    grupp_nr = groups_desert[i][0] - 1
    country = data.iloc[grupp_nr, 1]
    name = data.iloc[grupp_nr, 2]

    host_1_nr = starter_destination_group - 1
    destination_1 = data.iloc[host_1_nr, 4]
    address_1 = data.iloc[host_1_nr, 5]
    tel_1 = data.iloc[host_1_nr, 6]
    host_1_name = data.iloc[host_1_nr, 2]

    host_2_nr = main_destination_group - 1
    destination_2 = data.iloc[host_2_nr, 4]
    address_2 = data.iloc[host_2_nr, 5]
    tel_2 = data.iloc[host_2_nr, 6]
    host_2_name = data.iloc[host_2_nr, 2]

    guest_1 = groups_desert[i][1] - 1
    guest_2 = groups_desert[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, 7])
    special_diet21 = str(data.iloc[guest_2, 7])
    special_diet12 = str(data.iloc[guest_1, 8])
    special_diet22 = str(data.iloc[guest_2, 8])

    special_diet_line1 = 'Specialkost: KOST1 KOST2'
    special_diet_line1 = special_diet_line1.replace('KOST1', special_diet11)
    special_diet_line1 = special_diet_line1.replace('KOST2', special_diet12)

    special_diet_line2 = 'Specialkost: KOST1 KOST2'
    special_diet_line2 = special_diet_line2.replace('KOST1', special_diet21)
    special_diet_line2 = special_diet_line2.replace('KOST2', special_diet22)

    group_header = 'Grupp LAND, NAMN'
    group_header = group_header.replace('LAND', country)
    group_header = group_header.replace('NAMN', name)

    address_line = 'Adress: ADRESS'
    tel_line = 'Telefonnummer: NAME TEL'
    address_line = address_line.replace('ADRESS', address_1)
    tel_line = tel_line.replace('NAME', host_1_name)
    tel_line = tel_line.replace('TEL', tel_1)

    address_line2 = 'Adress: ADRESS'
    tel_line2 = 'Telefonnummer: NAME TEL'
    address_line2 = address_line2.replace('ADRESS', address_2)
    tel_line2 = tel_line2.replace('NAME', host_2_name)
    tel_line2 = tel_line2.replace('TEL', tel_2)

    body_text_1 = 'Resan tar nu sin början och det är nu dags för förrätt. Dags att vässa smaklökarna och bege er till ' \
                'er första destination OMRÅDE1..'
    body_text_2 = 'Redo för något exotiskt, det är nu dags för varmrätt. Ni ska nu ta cykeln vidare till OMRÅDE2, ' \
                'så drick upp glasen och tacka för förrätten. '
    body_text_1 = body_text_1.replace('OMRÅDE1', destination_1)
    body_text_2 = body_text_2.replace('OMRÅDE2', destination_2)

    doc2.add_heading(group_header, 2)

    doc2.add_heading('CykelfIEzten 2021', 2)
    doc2.add_paragraph(body_text_1)
    doc2.add_paragraph(address_line)
    doc2.add_paragraph(tel_line)
    doc2.add_paragraph(" ")

    doc2.add_heading('CykelfIEzten 2021', 2)
    doc2.add_paragraph(body_text_2)
    doc2.add_paragraph(address_line2)
    doc2.add_paragraph(tel_line2)
    doc2.add_paragraph(" ")

    doc2.add_heading('CykelfIEzten 2021', 2)
    doc2.add_paragraph('Resan närmar sig sitt slut och det är dags för kvällens grand finale. Skynda er hem och '
                       'välkomna gästerna. Efter det möts vi för gemensam utgång 22:00.')
    doc2.add_paragraph(" ")
    doc2.add_paragraph(special_diet_line1)
    doc2.add_paragraph(special_diet_line2)
    doc2.add_page_break()


def grouping(data):
    if len(data) % 3 == 1:
        diff = 1
        data.drop(len(data) - 1, axis=0, inplace=True)
        print("Grupperingen går ej ihop.", diff, "st grupp ej inkluderad. Gruppen motsvarar den sista i svarsfilen")
    if len(data) % 3 == 2:
        diff = 2
        data.drop([len(data) - 2, len(data) - 1], axis=0, inplace=True)
        print("Grupperingen går ej ihop.", diff, "st grupper ej inkluderad. Grupperna motsvarar de 2"
                                                 " sista i svarsfilen")

    groups_starter = []
    groups_main = []
    groups_desert = []
    for i in range(1, len(data), 3):
        groups_starter.append((i, i + 1, i + 2))

    for i in range(2, len(data), 3):
        if i + 4 == len(data):
            groups_main.append((i, i + 2, 3))
            groups_main.append((i + 3, 1, 6))
            break
        else:
            groups_main.append((i, i + 2, i + 7))

    for i in range(3, len(data), 3):
        if i + 3 == len(data):
            groups_desert.append((i, i + 1, 2))
            groups_desert.append((i + 3, 1, 5))
            break
        else:
            groups_desert.append((i, i + 1, i + 5))

    return groups_starter, groups_main, groups_desert


def desert_on_tviste(data):
    yo = list(range(2, 100, 3))
    yo1 = list(range(0, 100))
    for i in range(0, len(yo)):
        yo1.remove(yo[i])

    list_whit_rownumbers = []
    for i in range(0, len(data)):
        destination = data.iloc[i, 4]
        if destination == "Tvistevägen/Ålidhöjd" and i in yo1:
            list_whit_rownumbers.append(i)

    for i in range(0, len(data)):
        destination = data.iloc[i, 4]
        if destination != "Tvistevägen/Ålidhöjd" and i in yo:
            row_1 = data.iloc[i, :]
            try:
                index = list_whit_rownumbers[0]
                list_whit_rownumbers.pop(0)
                row_2 = data.iloc[index, :]
                data.iloc[i, :] = row_2
                data.iloc[index, :] = row_1
            except:
                break

    return data


def main():
    data = pd.read_excel(r'Förfestmarathon.xlsx', header=0)
    data = data.drop_duplicates(subset=["Telefonnummer till den ni ska vara hos (07x-xxxxxxx)"], keep='last', ignore_index=True)
    new_data = desert_on_tviste(data)

    groups_starter, groups_main, groups_desert = grouping(new_data)

    df = pd.DataFrame(new_data)
    df.to_excel("newfile.xlsx")

"""
    doc = docx.Document()
    doc1 = docx.Document()
    doc2 = docx.Document()
    for i in range(0, len(groups_starter)):
        starter_documentation(doc, groups_starter, groups_main, groups_desert, data, i)
        main_course_documentation(doc1, groups_starter, groups_main, groups_desert, data, i)
        desert_documentation(doc2, groups_starter, groups_main, groups_desert, data, i)
    doc.save('Förrätter.docx')
    doc1.save('Varmrätter.docx')
    doc2.save('Efterrätter.docx')
"""


if __name__ == '__main__':
    main()
