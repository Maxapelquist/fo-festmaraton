""""
    ==========================================
     Title:  Förfestmarathon
     Author: Max Apelquist
     Date:   8 Sep 2023
    ==========================================
"""

import pandas as pd
import docx


def summary(doc, groups_starter, groups_main, groups_desert, data, i):

    ## Förfest 1
    group_nr = groups_starter[i][0] - 1

    medlem1 = str(data.iloc[group_nr, 3])
    medlem2 = str(data.iloc[group_nr, 4])
    medlem3 = str(data.iloc[group_nr, 5])

    guest_1 = groups_starter[i][1] - 1
    guest_2 = groups_starter[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, 10])
    special_diet21 = str(data.iloc[guest_2, 10])

    special_diet12 = str(data.iloc[guest_1, 11])
    special_diet22 = str(data.iloc[guest_2, 11])

    special_diet13 = str(data.iloc[guest_1, 12])
    special_diet23 = str(data.iloc[guest_2, 12])

    special_diet14 = str(data.iloc[guest_1, 13])
    special_diet24 = str(data.iloc[guest_2, 13])

    special_kost_line1 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)
    special_kost_line1 = special_kost_line1.replace('KOST3', special_diet13)
    special_kost_line1 = special_kost_line1.replace('KOST4', special_diet14)

    special_kost_line2 = 'Specialkost: KOST1 , KOST2 , KOST3  , KOST4'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)
    special_kost_line2 = special_kost_line2.replace('KOST3', special_diet23)
    special_kost_line2 = special_kost_line2.replace('KOST4', special_diet24)


    group_header = 'Grupp: Namn1, Namn2, Namn3'
    group_header = group_header.replace('Namn1', medlem1)
    group_header = group_header.replace('Namn2', medlem2)
    group_header = group_header.replace('Namn3', medlem3)

    doc.add_heading(group_header, 2)
    doc.add_paragraph('Ni ska styra kvällens första förfest klockan 17:45. Ta hänsyn till följande specialkost: ')
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")

    ##Förfest 2
    group_nr = groups_main[i][0] - 1

    medlem1 = str(data.iloc[group_nr, 3])
    medlem2 = str(data.iloc[group_nr, 4])
    medlem3 = str(data.iloc[group_nr, 5])

    guest_1 = groups_main[i][1] - 1
    guest_2 = groups_main[i][2] - 1


    special_diet11 = str(data.iloc[guest_1, 9])
    special_diet21 = str(data.iloc[guest_2, 9])

    special_diet12 = str(data.iloc[guest_1, 10])
    special_diet22 = str(data.iloc[guest_2, 10])

    special_diet13 = str(data.iloc[guest_1, 11])
    special_diet23 = str(data.iloc[guest_2, 11])

    special_diet14 = str(data.iloc[guest_1, 12])
    special_diet24 = str(data.iloc[guest_2, 12])

    special_kost_line1 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)
    special_kost_line1 = special_kost_line1.replace('KOST3', special_diet13)
    special_kost_line1 = special_kost_line1.replace('KOST4', special_diet14)

    special_kost_line2 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)
    special_kost_line2 = special_kost_line2.replace('KOST3', special_diet23)
    special_kost_line2 = special_kost_line2.replace('KOST4', special_diet24)


    group_header = 'Grupp: Namn1, Namn2, Namn3'
    group_header = group_header.replace('Namn1', medlem1)
    group_header = group_header.replace('Namn2', medlem2)
    group_header = group_header.replace('Namn3', medlem3)

    doc.add_heading(group_header, 2)
    doc.add_paragraph('Ni ska styra kvällens andra förfest klockan 19:15. Ta hänsyn till följande specialkost: ')
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")

    ## Förfest 3

    group_nr = groups_desert[i][0] - 1

    medlem1 = str(data.iloc[group_nr, 3])
    medlem2 = str(data.iloc[group_nr, 4])
    medlem3 = str(data.iloc[group_nr, 5])

    guest_1 = groups_desert[i][1] - 1
    guest_2 = groups_desert[i][2] - 1

    special_diet11 = str(data.iloc[guest_1, 9])
    special_diet21 = str(data.iloc[guest_2, 9])

    special_diet12 = str(data.iloc[guest_1, 10])
    special_diet22 = str(data.iloc[guest_2, 10])

    special_diet13 = str(data.iloc[guest_1, 11])
    special_diet23 = str(data.iloc[guest_2, 11])

    special_diet14 = str(data.iloc[guest_1, 12])
    special_diet24 = str(data.iloc[guest_2, 12])

    special_kost_line1 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line1 = special_kost_line1.replace('KOST1', special_diet11)
    special_kost_line1 = special_kost_line1.replace('KOST2', special_diet12)
    special_kost_line1 = special_kost_line1.replace('KOST3', special_diet13)
    special_kost_line1 = special_kost_line1.replace('KOST4', special_diet14)

    special_kost_line2 = 'Specialkost: KOST1 , KOST2 , KOST3 , KOST4'
    special_kost_line2 = special_kost_line2.replace('KOST1', special_diet21)
    special_kost_line2 = special_kost_line2.replace('KOST2', special_diet22)
    special_kost_line2 = special_kost_line2.replace('KOST3', special_diet23)
    special_kost_line2 = special_kost_line2.replace('KOST4', special_diet24)

    group_header = 'Grupp: Namn1, Namn2, Namn3'
    group_header = group_header.replace('Namn1', medlem1)
    group_header = group_header.replace('Namn2', medlem2)
    group_header = group_header.replace('Namn3', medlem3)

    doc.add_heading(group_header, 2)
    doc.add_paragraph('Ni ska styra kvällens sista förfest klockan 20:45. Ta hänsyn till följande specialkost: ')
    doc.add_paragraph(special_kost_line1)
    doc.add_paragraph(special_kost_line2)
    doc.add_paragraph(" ")


def starter_documentation(doc, groups_starter, groups_main, groups_desert, data, i):

    starter_group_nr = groups_starter[i][0]
    for var in range(0, len(groups_main)):
        if starter_group_nr in groups_main[var]:
            main_destination_group = groups_main[var][0]

    for var1 in range(0, len(groups_desert)):
        if starter_group_nr in groups_desert[var1]:
            desert_destination_group = groups_desert[var1][0]

    group_nr = groups_starter[i][0] - 1
    country = data.iloc[group_nr, 2]

    host_1_nr = main_destination_group - 1
    destination_1 = str(data.iloc[host_1_nr, 7])
    address_1 = str(data.iloc[host_1_nr, 8])
    tel_1 = str(data.iloc[host_1_nr, 9])
    host_1_name = str(data.iloc[host_1_nr, 4])
    tema1 = str(data.iloc[host_1_nr, 2])


    host_2_nr = desert_destination_group - 1
    destination_2 = str(data.iloc[host_2_nr, 7])
    address_2 = str(data.iloc[host_2_nr, 8])
    tel_2 = str(data.iloc[host_2_nr, 9])
    host_2_name = str(data.iloc[host_2_nr, 4])
    tema2 = str(data.iloc[host_2_nr, 2])



    group_header = 'Förfestmarathon 2023 -  Grupp: LAND'
    group_header = group_header.replace('LAND', country)

    address_line = 'Adress: ADRESS'
    tema_line1 = 'Tema: TEMA'
    tel_line = 'Telefonnummer: NAME TEL'
    address_line = address_line.replace('ADRESS', address_1)
    tema_line1 = tema_line1.replace('TEMA', tema1)
    tel_line = tel_line.replace('NAME', host_1_name)
    tel_line = tel_line.replace('TEL', tel_1)

    address_line2 = 'Adress: ADRESS'
    tema_line2 = 'Tema: TEMA'
    tel_line2 = 'Telefonnummer: NAME TEL'
    address_line2 = address_line2.replace('ADRESS', address_2)
    tema_line2 = tema_line2.replace('TEMA', tema2)
    tel_line2 = tel_line2.replace('NAME', host_2_name)
    tel_line2 = tel_line2.replace('TEL', tel_2)



    body_text_1 = 'Har ni koll på klockan? Nästa förfest drar igång 19:15. Ni ska nu ta er vidare till OMRÅDE1,' \
                ' Drick upp glasen och tacka för ett bra krök.'
    body_text_2 = 'Resan närmar sig sitt slut och det är dags för kvällens sista förfest. Förfesten drar igång 20:45. Skynda er till OMRÅDE2 för att ' \
                'svalka era strupar innan gemensam utgång.'
    body_text_1 = body_text_1.replace('OMRÅDE1', destination_1)
    body_text_2 = body_text_2.replace('OMRÅDE2', destination_2)

    doc.add_heading(group_header, 1)
    doc.add_heading('Förfest 1', 2)
    doc.add_paragraph('Resan tar nu sin början och det är dags för er att styra förfest. Skynda er hem, gästerna anländer 17:45')
    doc.add_paragraph(" ")

    doc.add_heading('Förfest 2', 2)
    doc.add_paragraph(body_text_1)
    doc.add_paragraph(tema_line1)
    doc.add_paragraph(address_line)
    doc.add_paragraph(tel_line)
    doc.add_paragraph(" ")

    doc.add_heading('Förfest 3', 2)
    doc.add_paragraph(body_text_2)
    doc.add_paragraph(tema_line2)
    doc.add_paragraph(address_line2)
    doc.add_paragraph(tel_line2)
    doc.add_paragraph(" ")

    doc.add_heading("22:00 Location: Rouge Nightclub", 2)

    doc.add_page_break()


def main_course_documentation(doc1, groups_starter, groups_main, groups_desert, data, i):

    main_group_nr = groups_main[i][0]
    for var in range(0, len(groups_starter)):
        if main_group_nr in groups_starter[var]:
            starter_destination_group = groups_starter[var][0]

    for var1 in range(0, len(groups_desert)):
        if main_group_nr in groups_desert[var1]:
            desert_destination_group = groups_desert[var1][0]


    group_nr = groups_main[i][0] - 1
    country = data.iloc[group_nr, 2]

    host_1_nr = starter_destination_group - 1
    destination_1 = str(data.iloc[host_1_nr, 7])
    address_1 = str(data.iloc[host_1_nr, 8])
    tel_1 = str(data.iloc[host_1_nr, 9])
    host_1_name = str(data.iloc[host_1_nr, 4])
    tema1 = str(data.iloc[host_1_nr, 2])

    host_2_nr = desert_destination_group - 1
    destination_2 = str(data.iloc[host_2_nr, 7])
    address_2 = str(data.iloc[host_2_nr, 8])
    tel_2 = str(data.iloc[host_2_nr, 9])
    host_2_name = str(data.iloc[host_2_nr, 4])
    tema2 = str(data.iloc[host_2_nr, 2])



    group_header = 'Förfestmarathon 2023  -  Grupp LAND'
    group_header = group_header.replace('LAND', country)

    address_line = 'Adress: ADRESS'
    tel_line = 'Telefonnummer: NAME TEL'
    temaline_1 = 'Tema: TEMA'
    temaline_1 = temaline_1.replace('TEMA', tema1)
    address_line = address_line.replace('ADRESS', address_1)
    tel_line = tel_line.replace('NAME', host_1_name)
    tel_line = tel_line.replace('TEL', tel_1)

    address_line2 = 'Adress: ADRESS'
    tel_line2 = 'Telefonnummer: NAME TEL'
    temaline_2 = 'Tema: TEMA'
    temaline_2 = temaline_2.replace('TEMA', tema2)
    address_line2 = address_line2.replace('ADRESS', address_2)
    tel_line2 = tel_line2.replace('NAME', host_2_name)
    tel_line2 = tel_line2.replace('TEL', tel_2)



    body_text_1 = 'Resan tar nu sin början. Skynda er till OMRÅDE1 för kvällens första förfest. Festen drar igång 17:45'
    body_text_2 = 'Resan närmar sig sitt slut och det är dags för kvällens sista förfest. Förfesten drar igång 20:45. Skynda er till OMRÅDE2 för att ' \
                'svalka era strupar innan gemensam utgång klockan 22:00.'
    body_text_1 = body_text_1.replace('OMRÅDE1', destination_1)
    body_text_2 = body_text_2.replace('OMRÅDE2', destination_2)

    doc1.add_heading(group_header, 1)

    doc1.add_heading('Förfest 1', 2)
    doc1.add_paragraph(body_text_1)
    doc1.add_paragraph(temaline_1)
    doc1.add_paragraph(address_line)
    doc1.add_paragraph(tel_line)
    doc1.add_paragraph(" ")

    doc1.add_heading('Förfest 2', 2)
    doc1.add_paragraph('Det är dags för er att styra förfest. Skynda er hem, gästerna anländer 19:15')
    doc1.add_paragraph(" ")

    doc1.add_heading('Förfest 3', 2)
    doc1.add_paragraph(body_text_2)
    doc1.add_paragraph(temaline_2)
    doc1.add_paragraph(address_line2)
    doc1.add_paragraph(tel_line2)
    doc1.add_paragraph(" ")

    doc1.add_heading("22:00 Location: Rouge Nightclub", 2)

    doc1.add_page_break()


def desert_documentation(doc2, groups_starter, groups_main, groups_desert, data, i):

    desert_group_nr = groups_desert[i][0]
    for var in range(0, len(groups_starter)):
        if desert_group_nr in groups_starter[var]:
            starter_destination_group = groups_starter[var][0]

    for var1 in range(0, len(groups_main)):
        if desert_group_nr in groups_main[var1]:
            main_destination_group = groups_main[var1][0]

    grupp_nr = groups_desert[i][0] - 1
    country = data.iloc[grupp_nr, 2]

    host_1_nr = starter_destination_group - 1
    destination_1 = str(data.iloc[host_1_nr, 7])
    address_1 = str(data.iloc[host_1_nr, 8])
    tel_1 = str(data.iloc[host_1_nr, 9])
    host_1_name = str(data.iloc[host_1_nr, 4])
    tema1 = str(data.iloc[host_1_nr, 2])

    host_2_nr = main_destination_group - 1
    destination_2 = str(data.iloc[host_2_nr, 7])
    address_2 = str(data.iloc[host_2_nr, 8])
    tel_2 = str(data.iloc[host_2_nr, 9])
    host_2_name = str(data.iloc[host_2_nr, 4])
    tema2 = str(data.iloc[host_2_nr, 2])

    group_header = 'Förfestmarathon 2023 - Grupp: LAND'
    group_header = group_header.replace('LAND', country)

    address_line = 'Adress: ADRESS'
    tel_line = 'Telefonnummer: NAME TEL'
    temaline_1 = 'Tema: TEMA'
    temaline_1 = temaline_1.replace('TEMA', tema1)
    address_line = address_line.replace('ADRESS', address_1)
    tel_line = tel_line.replace('NAME', host_1_name)
    tel_line = tel_line.replace('TEL', tel_1)

    address_line2 = 'Adress: ADRESS'
    tel_line2 = 'Telefonnummer: NAME TEL'
    temaline_2 = 'Tema: TEMA'
    temaline_2 = temaline_2.replace('TEMA', tema2)
    address_line2 = address_line2.replace('ADRESS', address_2)
    tel_line2 = tel_line2.replace('NAME', host_2_name)
    tel_line2 = tel_line2.replace('TEL', tel_2)


    body_text_1 = 'Resan tar nu sin början. Skynda er till OMRÅDE1 för kvällens första förfest. Festen drar igång 17:45'

    body_text_2 = 'Har ni koll på klockan? Nästa förfest drar igång 19:15. Ni ska nu ta er vidare till OMRÅDE2,' \
                ' Drick upp glasen och tacka för ett bra krök.'
    body_text_1 = body_text_1.replace('OMRÅDE1', destination_1)
    body_text_2 = body_text_2.replace('OMRÅDE2', destination_2)

    doc2.add_heading(group_header, 1)

    doc2.add_heading('Förfest 1', 2)
    doc2.add_paragraph(body_text_1)
    doc2.add_paragraph(temaline_1)
    doc2.add_paragraph(address_line)
    doc2.add_paragraph(tel_line)
    doc2.add_paragraph(" ")

    doc2.add_heading('Förfest 2', 2)
    doc2.add_paragraph(body_text_2)
    doc2.add_paragraph(temaline_2)
    doc2.add_paragraph(address_line2)
    doc2.add_paragraph(tel_line2)
    doc2.add_paragraph(" ")

    doc2.add_heading('Förfest 3', 2)
    doc2.add_paragraph('Resan närmar sig sitt slut och det är dags för kvällens grand finale. Skynda er hem och '
                       'välkomna gästerna. Efter det möts vi för gemensam utgång 22:00.')

    doc2.add_paragraph(" ")
    doc2.add_heading("22:00 Location: Rouge Nightclub", 2)
    doc2.add_page_break()


def grouping(data):
    if len(data) % 3 == 1:
        diff = 1
        data.drop(len(data) - 1, axis=0, inplace=True)
        print("Grupperingen går ej ihop.", diff, "st grupp ej inkluderad. Gruppen motsvarar den sista i svarsfilen")
    elif len(data) % 3 == 2:
        diff = 2
        data.drop([len(data) - 2, len(data) - 1], axis=0, inplace=True)
        print("Grupperingen går ej ihop.", diff, "st grupper ej inkluderad. Grupperna motsvarar de 2"
                                                 " sista i svarsfilen")
    else:
        print("Grupperingen går ihop, antal grupper: ", len(data) )

    groups_starter = []
    groups_main = []
    groups_desert = []
    for i in range(1, len(data), 3):
        groups_starter.append((i, i + 1, i + 2))

    for i in range(2, len(data), 3):
        if i + 4 == len(data):
            groups_main.append((i, i + 2, 3))
            groups_main.append((i + 3, 1, 6))
            break
        else:
            groups_main.append((i, i + 2, i + 7))

    for i in range(3, len(data), 3):
        if i + 3 == len(data):
            groups_desert.append((i, i + 1, 2))
            groups_desert.append((i + 3, 1, 5))
            break
        else:
            groups_desert.append((i, i + 1, i + 5))

    return groups_starter, groups_main, groups_desert


def last_on_tviste(data, column_name):
    yo = list(range(2, 100, 3))
    yo1 = list(range(0, 100))
    for i in range(0, len(yo)):
        yo1.remove(yo[i])

    list_whit_rownumbers = []
    for i in range(0, len(data)):
        destination = data.iloc[i, data.columns.get_loc(column_name)]
        if destination == "Tvistevägen/Ålidhöjd" and i in yo1:
            list_whit_rownumbers.append(i)

    for i in range(0, len(data)):
        destination = data.iloc[i, data.columns.get_loc(column_name)]
        if destination != "Tvistevägen/Ålidhöjd" and i in yo:
            row_1 = data.iloc[i, :]
            try:
                index = list_whit_rownumbers[0]
                list_whit_rownumbers.pop(0)
                row_2 = data.iloc[index, :]
                data.iloc[i, :] = row_2
                data.iloc[index, :] = row_1
            except:
                break

    return data


def sort_data(data):
    # Skapa en tom dataframe för att lagra den sorterade data
    new_data = pd.DataFrame(columns=data.columns)

    # Sortera data baserat på kolumnen 'Vilket program pluggar ni?'
    sorted_data = data.sort_values(by=['Viket program pluggar ni?'])

    # Dela upp de sorterade raderna i tre olika dataframes baserat på vilket program som anges i kolumnen
    df1 = sorted_data[sorted_data['Viket program pluggar ni?'] == 'Industriell Ekonomi']
    df2 = sorted_data[sorted_data['Viket program pluggar ni?'] == 'Arkitekt / Design']
    df3 = sorted_data[sorted_data['Viket program pluggar ni?'] == 'Jurist']

    # Hitta längden på den kortaste av de tre dataframesen och använd den för att säkerställa att inga rader går förlorade
    min_length = min(len(df1), len(df2), len(df3))

    # Iterera igenom de tre dataframesen samtidigt och lägg till varannan rad i den nya dataramen
    for i in range(min_length):
        new_data = new_data._append(df1.iloc[i])
        new_data = new_data._append(df2.iloc[i])
        new_data = new_data._append(df3.iloc[i])

    # Lägg till eventuella kvarvarande rader från de tre ursprungliga dataframesen
    new_data = new_data._append(df1[min_length:])
    new_data = new_data._append(df2[min_length:])
    new_data = new_data._append(df3[min_length:])

    # Återställ index i den sorterade dataramen
    new_data.reset_index(drop=True, inplace=True)

    return new_data


def main():
    data = pd.read_excel(r'Förfestmarathon 2024.xlsx', header=0)

    #Inkluderades något år för att vissa personer anmälde sig flera gånger. Om samma telefonummer upkommer flera gånger sparas bara den sista anmälan
    data = data.drop_duplicates(subset=["Telefonnummer till den ni ska vara hos (07x-xxxxxxx)"], keep='last', ignore_index=True)

    #Funktion sorterar datat för att olika program skall hamna på olika förfester
    data1 = sort_data(data)

    #Funktion sorterar datat för att efterrätter skall hamna på tviste
    data2 = last_on_tviste(data1, "Vilket område kommer ni vara på?")

    #utför grupperingen
    groups_starter, groups_main, groups_desert = grouping(data2)

    #Onödigt, men sparar filen för att kunna kontrollera visuellt hur grupperingarna blev.
    df = pd.DataFrame(data2)
    df.to_excel("newfile1.xlsx")

    #Här kan justeringar göras för hand, typ denna grupp skall träffa dessa osv osv
    #Juster i filen newfile1.xlsx.
    #Spara somm ny fil "final.xlsx" och kör denna rad: finaldata = pd.read_excel(r'final.xlsx', header=0)
    #Kom ihåg att kommentera ut raden under isåfall "finaldata = pd.read_excel(r'newfile1.xlsx', header=0)"

    #läser in datat igen, OBS: om ändringar ej gjorts
    finaldata = pd.read_excel(r'newfile1.xlsx', header=0)


    # Skapa summary dokument för publicering innan event. Innehållande vilken grupp som har vilken rätt och vilken specialkost de bör ta hänsyn till
    doc = docx.Document()
    for k in range(0, len(groups_starter)):
        summary(doc, groups_starter, groups_main, groups_desert, finaldata, k)
    doc.save('outputs/Summary.docx')


    #Skapa dokument "breven"
    doc1 = docx.Document()
    doc2 = docx.Document()
    doc3 = docx.Document()
    for i in range(0, len(groups_starter)):
        starter_documentation(doc1, groups_starter, groups_main, groups_desert, finaldata, i)
        main_course_documentation(doc2, groups_starter, groups_main, groups_desert, finaldata, i)
        desert_documentation(doc3, groups_starter, groups_main, groups_desert, finaldata, i)
    doc1.save('outputs/Förrätter.docx')
    doc2.save('outputs/Varmrätter.docx')
    doc3.save('outputs/Efterrätter.docx')



if __name__ == '__main__':
    main()